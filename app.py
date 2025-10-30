import base64
import os
import platform
import subprocess
import traceback
import uuid
from datetime import datetime
from io import BytesIO

from typing import Any, Optional, cast

from docx import Document
from docx.shared import Mm
try:
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - handled at runtime
    WD_ROW_HEIGHT_RULE = None  # type: ignore[assignment]
    WD_ALIGN_VERTICAL = None  # type: ignore[assignment]
from flask import Flask, jsonify, request, send_file
from flasgger import Swagger
from num2words import num2words

try:
    import qrcode  # type: ignore[import-not-found]
    from qrcode.constants import ERROR_CORRECT_M  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - handled at runtime
    qrcode = None  # type: ignore[assignment]
    ERROR_CORRECT_M = None  # type: ignore[assignment]

try:
    from PIL import Image  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - handled at runtime
    Image = None  # type: ignore[assignment]
import zipfile

TEMPLATE_PATH = "./Templates/LeadsForce_v0.docx"
OUTPUT_DIR = "./output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

PLACEHOLDERS = [
    "ID", "INVOICE_DATE", "CUSTOMER", "PRODUCT", "SUM", "AMOUNT_IN_WORDS",
    "DEAL", "SERVICE", "CITY", "LEAD_SUM", "LEAD_COST", "REVENUE", "PRICE",
    "EMAIL", "PHONE", "NAME", "INN", "COMPANYNAME", "PAYMENT_QR_BASE64",
    "PAYMENT_QR_PAYLOAD"
]

app = Flask(__name__)

SWAGGER_PARAMETERS = {
    "price": {
        "name": "price",
        "in": "query",
        "description": "Стоимость в рублях (например, 12345.67)",
        "schema": {"type": "string"}
    },
    "price_text": {
        "name": "price_text",
        "in": "query",
        "description": "Сумма прописью, если требуется переопределить автоматическую генерацию",
        "schema": {"type": "string"}
    },
    "bill_date": {
        "name": "bill_date",
        "in": "query",
        "description": "Дата счёта (формат дд.мм.гггг)",
        "schema": {"type": "string"}
    },
    "invoiceDate": {
        "name": "invoiceDate",
        "in": "query",
        "description": "Альтернативный параметр даты счёта",
        "schema": {"type": "string"}
    },
    "deal": {
        "name": "deal",
        "in": "query",
        "description": "Номер сделки/счёта",
        "schema": {"type": "string"}
    },
    "service": {
        "name": "service",
        "in": "query",
        "description": "Название услуги",
        "schema": {"type": "string"}
    },
    "city": {
        "name": "city",
        "in": "query",
        "description": "Город клиента",
        "schema": {"type": "string"}
    },
    "lead_sum": {
        "name": "lead_sum",
        "in": "query",
        "description": "Общая сумма лида",
        "schema": {"type": "string"}
    },
    "lead_cost": {
        "name": "lead_cost",
        "in": "query",
        "description": "Стоимость лида",
        "schema": {"type": "string"}
    },
    "revenue": {
        "name": "revenue",
        "in": "query",
        "description": "Выручка",
        "schema": {"type": "string"}
    },
    "email": {
        "name": "email",
        "in": "query",
        "description": "Email клиента",
        "schema": {"type": "string", "format": "email"}
    },
    "phone": {
        "name": "phone",
        "in": "query",
        "description": "Телефон клиента",
        "schema": {"type": "string"}
    },
    "name": {
        "name": "name",
        "in": "query",
        "description": "Имя клиента",
        "schema": {"type": "string"}
    },
    "inn": {
        "name": "inn",
        "in": "query",
        "description": "ИНН клиента",
        "schema": {"type": "string"}
    },
    "companyName": {
        "name": "companyName",
        "in": "query",
        "description": "Название компании клиента",
        "schema": {"type": "string"}
    },
    "qr_sum": {
        "name": "qr_sum",
        "in": "query",
        "description": "Сумма платежа в копейках для QR",
        "schema": {"type": "string"}
    },
    "qr_purpose": {
        "name": "qr_purpose",
        "in": "query",
        "description": "Назначение платежа для QR",
        "schema": {"type": "string"}
    },
    "qr_width_mm": {
        "name": "qr_width_mm",
        "in": "query",
        "description": "Ширина QR в миллиметрах при вставке в шаблон",
        "schema": {"type": "number"}
    },
    "qr_name": {
        "name": "qr_name",
        "in": "query",
        "description": "Получатель платежа",
        "schema": {"type": "string"}
    },
    "qr_personal_account": {
        "name": "qr_personal_account",
        "in": "query",
        "description": "Расчётный счёт получателя",
        "schema": {"type": "string"}
    },
    "qr_bank_name": {
        "name": "qr_bank_name",
        "in": "query",
        "description": "Название банка",
        "schema": {"type": "string"}
    },
    "qr_bic": {
        "name": "qr_bic",
        "in": "query",
        "description": "БИК банка",
        "schema": {"type": "string"}
    },
    "qr_correspondent_account": {
        "name": "qr_correspondent_account",
        "in": "query",
        "description": "Корреспондентский счёт",
        "schema": {"type": "string"}
    },
    "qr_inn": {
        "name": "qr_inn",
        "in": "query",
        "description": "ИНН получателя",
        "schema": {"type": "string"}
    },
    "qr_kpp": {
        "name": "qr_kpp",
        "in": "query",
        "description": "КПП получателя",
        "schema": {"type": "string"}
    },
    "qr_payer_address": {
        "name": "qr_payer_address",
        "in": "query",
        "description": "Адрес плательщика",
        "schema": {"type": "string"}
    }
}

swagger_template = {
    "swagger": "2.0",
    "info": {
        "title": "LeadForce Document Generator",
        "description": "API для генерации документов и банковских QR-кодов",
        "version": "1.0.0"
    },
    "basePath": "/",
    "parameters": SWAGGER_PARAMETERS,
}

swagger_config = {
    "headers": [],
    "specs": [
        {
            "endpoint": "swagger",
            "route": "/openapi.json",
            "rule_filter": lambda rule: True,
            "model_filter": lambda tag: True,
        }
    ],
    "static_url_path": "/flasgger_static",
    "swagger_ui": True,
    "specs_route": "/apidocs/"
}

swagger = Swagger(app, template=swagger_template, config=swagger_config)

DEFAULT_PAYMENT_DETAILS = {
    "Name": "ИП Абакумова Наталья Александровна",
    "PersonalAcc": "40802810200006322048",
    "BankName": "АО «Тинькофф Банк»",
    "BIC": "044525974",
    "CorrespAcc": "30101810145250000974",
    "PayeeINN": "720206359451",
    "Purpose": "Оплата по счету №{{ID}}"
}

QR_QUERY_MAP = {
    "qr_name": "Name",
    "qr_personal_account": "PersonalAcc",
    "qr_bank_name": "BankName",
    "qr_bic": "BIC",
    "qr_correspondent_account": "CorrespAcc",
    "qr_inn": "PayeeINN",
    "qr_kpp": "PayeeKPP",
    "qr_payer_address": "PayerAddress"
}

QR_CODE_PLACEHOLDER = "{{QR_CODE}}"

def fill_template_xml(template_path: str, replacements: dict, output_path: str):
    with zipfile.ZipFile(template_path, 'r') as zin:
        with zipfile.ZipFile(output_path, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    xml = data.decode('utf-8')
                    for key, value in replacements.items():
                        xml = xml.replace(f'{{{{{key}}}}}', value)
                    data = xml.encode('utf-8')
                zout.writestr(item, data)

def convert_to_pdf(input_docx: str, output_dir: str):
    if platform.system() == "Windows":
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore

        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            abs_input = os.path.abspath(input_docx).replace("/", "\\")
            doc = word.Documents.Open(abs_input)
            output_pdf = os.path.splitext(abs_input)[0] + ".pdf"
            doc.SaveAs(output_pdf, FileFormat=17)
            doc.Close(False)
            word.Quit()
            return output_pdf
        finally:
            pythoncom.CoUninitialize()
    else:
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf",
            "--outdir", output_dir, input_docx
        ], check=True)
        return os.path.splitext(input_docx)[0] + ".pdf"

def zip_single_file(file_path, arcname):
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as zf:
        zf.write(file_path, arcname=arcname)
    buffer.seek(0)
    return buffer

def zip_files(file_mappings):
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as zf:
        for path, arcname in file_mappings:
            if path and os.path.exists(path):
                zf.write(path, arcname=arcname)
    buffer.seek(0)
    return buffer


def encode_file_to_base64(file_path: str) -> str:
    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode("ascii")


PAYMENT_QR_FIELDS_ORDER = [
    "Name",
    "PersonalAcc",
    "BankName",
    "BIC",
    "CorrespAcc",
    "PayeeINN",
    "PayeeKPP",
    "PayerAddress",
    "Sum",
    "Purpose",
]


DEFAULT_QR_WIDTH_MM = 35

# Дополнительный запас, который мы оставляем внутри ячейки таблицы при вставке QR.
# На Linux LibreOffice при конвертации DOCX -> PDF заметно сильнее подрезает
# изображения, если они вплотную подходят к границам ячейки, поэтому держим
# небольшой фиксированный и относительный зазоры.
QR_CELL_MARGIN_MM = 2
QR_CELL_MARGIN_RATIO = 0.1

# На Linux при конвертации в PDF LibreOffice иногда обрезает саму картинку даже при
# соблюдении отступов в таблице, если у PNG слишком тонкая белая рамка. Поэтому мы
# принудительно добавляем запас вокруг QR-кода при сохранении файла.
QR_IMAGE_PADDING_PX = 24
QR_IMAGE_PADDING_RATIO = 0.12


def build_payment_qr_payload(details: dict) -> str:
    parts = ["ST00012"]
    used_keys = set()

    for field in PAYMENT_QR_FIELDS_ORDER:
        value = (details.get(field) or "").strip()
        if value:
            parts.append(f"{field}={value}")
            used_keys.add(field)

    for key, value in details.items():
        if key in used_keys:
            continue
        value = (value or "").strip()
        if value:
            parts.append(f"{key}={value}")

    return "|".join(parts)


def _require_qr_dependencies() -> Optional[str]:
    missing = []
    if qrcode is None:
        missing.append("qrcode")
    if Image is None:
        missing.append("Pillow")
    if missing:
        return ", ".join(missing)
    return None


def generate_payment_qr_image(details: dict, file_id: str) -> tuple[str, str]:
    missing = _require_qr_dependencies()
    if missing:
        raise RuntimeError(
            "Для генерации QR-кода необходимо установить зависимости: "
            f"{missing}. Выполните 'pip install -r requirements.txt'."
        )

    payload = build_payment_qr_payload(details)
    if len(payload) <= len("ST00012"):
        return "", ""

    qr_path = os.path.join(OUTPUT_DIR, f"{file_id}_qr.png")
    qr_module = cast(Any, qrcode)
    error_correction = cast(int, ERROR_CORRECT_M)
    qr = qr_module.QRCode(error_correction=error_correction, box_size=10, border=4)
    qr.add_data(payload)
    qr.make(fit=True)
    qr_image = qr.make_image(fill_color="black", back_color="white")
    pil_image = qr_image.get_image() if hasattr(qr_image, "get_image") else qr_image
    if not hasattr(pil_image, "save"):
        raise TypeError("Объект QR-кода не поддерживает сохранение в файл")
    pil_image.save(qr_path, format="PNG")
    _ensure_qr_image_padding(qr_path)

    return payload, qr_path


def _replace_paragraph_with_image(paragraph, image_path: str, width_mm: float):
    while paragraph.runs:
        paragraph._element.remove(paragraph.runs[0]._r)
    run = paragraph.add_run()
    run.add_picture(image_path, width=Mm(width_mm))


def _replace_in_paragraphs(paragraphs, placeholder: str, image_path: str, width_mm: float) -> bool:
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            _replace_paragraph_with_image(paragraph, image_path, width_mm)
            return True
    return False


def _ensure_qr_image_padding(image_path: str) -> None:
    """Добавляет белую рамку вокруг QR-кода, чтобы избежать обрезания при экспорте."""
    if Image is None:
        return

    try:
        with Image.open(image_path) as img:
            qr_image = img.convert("RGB")
            min_side = min(qr_image.size)
            padding = max(int(min_side * QR_IMAGE_PADDING_RATIO), QR_IMAGE_PADDING_PX)
            if padding <= 0:
                return

            new_size = (qr_image.width + padding * 2, qr_image.height + padding * 2)
            padded = Image.new("RGB", new_size, "white")
            padded.paste(qr_image, (padding, padding))
            padded.save(image_path, format="PNG")
    except Exception:
        traceback.print_exc()


def _apply_qr_margin(limit_mm: float) -> float:
    """Возвращает максимально допустимую ширину с учётом фиксированного и относительного зазоров."""
    if not limit_mm:
        return 0

    margin = max(QR_CELL_MARGIN_MM, limit_mm * QR_CELL_MARGIN_RATIO)
    return max(limit_mm - margin, 5)


def _ensure_cell_can_fit_image(row, cell, image_width_mm: float) -> None:
    """Настраивает параметры строки и ячейки таблицы, чтобы QR полностью уместился."""
    try:
        margin = max(QR_CELL_MARGIN_MM, image_width_mm * QR_CELL_MARGIN_RATIO)
        required_height_mm = image_width_mm + margin

        if WD_ROW_HEIGHT_RULE is not None:
            try:
                if getattr(row, "height_rule", None) == WD_ROW_HEIGHT_RULE.EXACTLY:
                    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            except Exception:
                traceback.print_exc()

        current_height_mm = getattr(getattr(row, "height", None), "mm", None)
        if current_height_mm and current_height_mm < required_height_mm:
            try:
                row.height = Mm(required_height_mm)
                if WD_ROW_HEIGHT_RULE is not None:
                    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            except Exception:
                traceback.print_exc()

        if WD_ALIGN_VERTICAL is not None and cell is not None:
            try:
                if cell.vertical_alignment is None:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                traceback.print_exc()
    except Exception:
        traceback.print_exc()


def _paragraph_has_placeholder(paragraph) -> bool:
    if QR_CODE_PLACEHOLDER in getattr(paragraph, "text", ""):
        return True

    try:
        return QR_CODE_PLACEHOLDER in ''.join(run.text for run in getattr(paragraph, "runs", []))
    except Exception:
        return False


def _clamp_width_to_cell(width_mm: float, row, cell) -> float:
    limits = []

    row_height = getattr(row.height, "mm", None)
    if row_height:
        limits.append(_apply_qr_margin(row_height))

    cell_width_attr = getattr(cell, "width", None)
    cell_width_mm = getattr(cell_width_attr, "mm", None) if cell_width_attr else None
    if cell_width_mm:
        limits.append(_apply_qr_margin(cell_width_mm))

    # Иногда python-docx не устанавливает ширину ячейки, но задаёт её в tcW (twips).
    if not limits:
        tc_pr = getattr(getattr(cell, "_tc", None), "tcPr", None)
        tc_w = getattr(tc_pr, "tcW", None) if tc_pr is not None else None
        width_twips = getattr(tc_w, "w", None) if tc_w is not None else None
        try:
            width_twips_int = int(width_twips) if width_twips is not None else None
        except (TypeError, ValueError):
            width_twips_int = None
        if width_twips_int:
            width_mm = width_twips_int * 25.4 / 1440
            limits.append(_apply_qr_margin(width_mm))

    if not limits:
        return min(width_mm, DEFAULT_QR_WIDTH_MM)

    safe_limit = min(limit for limit in limits if limit)
    return min(width_mm, safe_limit)


def insert_qr_code_into_document(docx_path: str, qr_image_path: str, width_mm: float) -> bool:
    document = Document(docx_path)

    if _replace_in_paragraphs(document.paragraphs, QR_CODE_PLACEHOLDER, qr_image_path, width_mm):
        document.save(docx_path)
        return True

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if any(_paragraph_has_placeholder(paragraph) for paragraph in cell.paragraphs):
                    effective_width = _clamp_width_to_cell(width_mm, row, cell)
                    _ensure_cell_can_fit_image(row, cell, effective_width)
                    if _replace_in_paragraphs(cell.paragraphs, QR_CODE_PLACEHOLDER, qr_image_path, effective_width):
                        document.save(docx_path)
                        return True

    return False

MONTHS_RU = {
    '01': 'января', '02': 'февраля', '03': 'марта',
    '04': 'апреля', '05': 'мая', '06': 'июня',
    '07': 'июля', '08': 'августа', '09': 'сентября',
    '10': 'октября', '11': 'ноября', '12': 'декабря'
}


def parse_sum_to_kopecks(price_str: str) -> str:
    normalized = (price_str or "").replace(" ", "").replace(",", ".")
    if not normalized:
        return ""

    try:
        amount = float(normalized)
    except ValueError:
        return ""

    return str(int(round(amount * 100)))


def get_qr_width_mm(args) -> float:
    value = (args.get("qr_width_mm", "") or "").strip()
    if not value:
        return DEFAULT_QR_WIDTH_MM

    try:
        width = float(value.replace(",", "."))
        if width <= 0:
            return DEFAULT_QR_WIDTH_MM
        return width
    except ValueError:
        return DEFAULT_QR_WIDTH_MM


def get_payment_details(args, replacements: dict) -> dict:
    details = DEFAULT_PAYMENT_DETAILS.copy()

    for query_param, payload_key in QR_QUERY_MAP.items():
        value = (args.get(query_param, "") or "").strip()
        if value:
            details[payload_key] = value

    invoice_id = replacements.get("ID", "")

    sum_override = (args.get("qr_sum", "") or "").strip()
    if sum_override:
        details["Sum"] = sum_override
    else:
        auto_sum = parse_sum_to_kopecks(replacements.get("SUM", ""))
        if auto_sum:
            details["Sum"] = auto_sum

    purpose_override = (args.get("qr_purpose", "") or "").strip()
    if purpose_override:
        details["Purpose"] = purpose_override
    else:
        purpose = details.get("Purpose", "")
        if purpose and "{{ID}}" in purpose:
            details["Purpose"] = purpose.replace("{{ID}}", invoice_id)
        elif not purpose and invoice_id:
            details["Purpose"] = f"Оплата по счету №{invoice_id}"

    return details


def format_invoice_date(date_str):
    try:
        day, month, year = date_str.strip().split('.')
        return f"{int(day)} {MONTHS_RU[month]} {year} г."
    except:
        return date_str  # fallback если не смогли

def get_replacements():
    args = request.args

    price_str = args.get("price", "").replace(",", ".").strip()
    price_text = args.get("price_text", "").strip()

    # Обработка даты
    bill_date_raw = args.get("bill_date", "").strip()
    bill_date = bill_date_raw.split()[0] if bill_date_raw and " " in bill_date_raw else ""
    invoice_date = bill_date or args.get("invoiceDate", "").strip()
    if not invoice_date:
        invoice_date = datetime.today().strftime('%d.%m.%Y')

    # Генерация суммы прописью
    try:
        if not price_text:
            price_float = float(price_str)
            rub = int(price_float)
            kop = int(round((price_float - rub) * 100))
            amount_in_words = f"{num2words(rub, lang='ru').capitalize()} рублей {kop:02d} копеек"
        else:
            amount_in_words = price_text
    except:
        amount_in_words = ""

    # Заказчик
    customer_parts = [
        args.get("name", "").strip(),
        args.get("phone", "").strip(),
        args.get("email", "").strip(),
        args.get("inn", "").strip(),
        args.get("companyName", "").strip()
    ]
    customer = ", ".join(filter(None, customer_parts))

    # Товар
    product_service = args.get("service", "").strip()
    product = f"Система привлечения клиентов / {product_service}" if product_service else "Система привлечения клиентов"

    return {
        "ID": args.get("deal", str(uuid.uuid4())[:8]),
        "INVOICE_DATE": format_invoice_date(invoice_date),
        "CUSTOMER": customer,
        "PRODUCT": product,
        "SUM": price_str,
        "AMOUNT_IN_WORDS": amount_in_words,
        "DEAL": args.get("deal", ""),
        "SERVICE": args.get("service", ""),
        "CITY": args.get("city", ""),
        "LEAD_SUM": args.get("lead_sum", ""),
        "LEAD_COST": args.get("lead_cost", ""),
        "REVENUE": args.get("revenue", ""),
        "PRICE": price_str,
        "EMAIL": args.get("email", ""),
        "PHONE": args.get("phone", ""),
        "NAME": args.get("name", ""),
        "INN": args.get("inn", ""),
        "COMPANYNAME": args.get("companyName", ""),
        "PAYMENT_QR_BASE64": "",
        "PAYMENT_QR_PAYLOAD": ""
    }


def prepare_generation_inputs():
    replacements = get_replacements()
    payment_details = get_payment_details(request.args, replacements)
    qr_width_mm = get_qr_width_mm(request.args)
    return replacements, payment_details, qr_width_mm


def build_doc(replacements: dict, payment_details: dict, qr_width_mm: float):
    file_id = str(uuid.uuid4())
    docx_path = os.path.join(OUTPUT_DIR, f"{file_id}.docx")

    replacements_for_template = dict(replacements)
    qr_payload = ""
    qr_path = ""

    try:
        qr_payload, qr_path = generate_payment_qr_image(payment_details, file_id)
    except Exception as qr_error:
        traceback.print_exc()
        replacements_for_template["PAYMENT_QR_PAYLOAD"] = str(qr_error)
        replacements_for_template["PAYMENT_QR_BASE64"] = ""

    if qr_payload and qr_path and os.path.exists(qr_path):
        try:
            replacements_for_template["PAYMENT_QR_PAYLOAD"] = qr_payload
            replacements_for_template["PAYMENT_QR_BASE64"] = encode_file_to_base64(qr_path)
        except Exception:
            traceback.print_exc()

    fill_template_xml(TEMPLATE_PATH, replacements_for_template, docx_path)

    if qr_payload and qr_path and os.path.exists(qr_path):
        try:
            insert_qr_code_into_document(docx_path, qr_path, qr_width_mm)
        except Exception:
            traceback.print_exc()

    pdf_path = convert_to_pdf(docx_path, OUTPUT_DIR)
    return docx_path, pdf_path, qr_path

def _build_service_description() -> dict:
    return {
        "message": "LeadForce Document Generator", 
        "endpoints": {
            "pdf": "/Document/GetPdf",
            "docx": "/Document/GetDocx",
            "zip_pdf": "/Document/GetPdfZip",
            "zip_docx": "/Document/GetDocxZip",
            "zip_all": "/Document/GetAllZip",
            "qr_png": "/Document/GetPaymentQr"
        },
        "docs": "Отправьте GET-запрос на любой endpoint, передав параметры сделки в query string."
    }


@app.route("/")
def index():
    """Описание сервиса
    ---
    tags:
      - Service
    responses:
      200:
        description: JSON со списком доступных маршрутов
    """
    return jsonify(_build_service_description())


@app.route("/docs")
def docs():
    """Дублирующий endpoint описания сервиса
    ---
    tags:
      - Service
    responses:
      200:
        description: JSON со списком доступных маршрутов
    """
    return jsonify(_build_service_description())


@app.route("/favicon.ico")
def favicon():
    return ("", 204)


@app.route("/Document/GetPdf")
def get_pdf():
    """Получить PDF с заполненным шаблоном
    ---
    tags:
      - Documents
    produces:
      - application/pdf
    parameters:
      - $ref: '#/parameters/price'
      - $ref: '#/parameters/price_text'
      - $ref: '#/parameters/bill_date'
      - $ref: '#/parameters/invoiceDate'
      - $ref: '#/parameters/deal'
      - $ref: '#/parameters/service'
      - $ref: '#/parameters/city'
      - $ref: '#/parameters/lead_sum'
      - $ref: '#/parameters/lead_cost'
      - $ref: '#/parameters/revenue'
      - $ref: '#/parameters/email'
      - $ref: '#/parameters/phone'
      - $ref: '#/parameters/name'
      - $ref: '#/parameters/inn'
      - $ref: '#/parameters/companyName'
      - $ref: '#/parameters/qr_sum'
      - $ref: '#/parameters/qr_purpose'
      - $ref: '#/parameters/qr_width_mm'
      - $ref: '#/parameters/qr_name'
      - $ref: '#/parameters/qr_personal_account'
      - $ref: '#/parameters/qr_bank_name'
      - $ref: '#/parameters/qr_bic'
      - $ref: '#/parameters/qr_correspondent_account'
      - $ref: '#/parameters/qr_inn'
      - $ref: '#/parameters/qr_kpp'
      - $ref: '#/parameters/qr_payer_address'
    responses:
      200:
        description: PDF файл с заполненными данными
        content:
          application/pdf:
            schema:
              type: string
              format: binary
      500:
        description: Ошибка генерации документа
    """
    try:
        replacements, payment_details, qr_width_mm = prepare_generation_inputs()
        _, pdf_path, _ = build_doc(replacements, payment_details, qr_width_mm)
        return send_file(
            pdf_path,
            download_name="document.pdf",
            mimetype="application/pdf",
            as_attachment=True,
        )
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/Document/GetDocx")
def get_docx():
    """Получить DOCX с заполненным шаблоном
    ---
    tags:
      - Documents
    produces:
      - application/vnd.openxmlformats-officedocument.wordprocessingml.document
    parameters:
      - $ref: '#/parameters/price'
      - $ref: '#/parameters/price_text'
      - $ref: '#/parameters/bill_date'
      - $ref: '#/parameters/invoiceDate'
      - $ref: '#/parameters/deal'
      - $ref: '#/parameters/service'
      - $ref: '#/parameters/city'
      - $ref: '#/parameters/lead_sum'
      - $ref: '#/parameters/lead_cost'
      - $ref: '#/parameters/revenue'
      - $ref: '#/parameters/email'
      - $ref: '#/parameters/phone'
      - $ref: '#/parameters/name'
      - $ref: '#/parameters/inn'
      - $ref: '#/parameters/companyName'
      - $ref: '#/parameters/qr_sum'
      - $ref: '#/parameters/qr_purpose'
      - $ref: '#/parameters/qr_width_mm'
      - $ref: '#/parameters/qr_name'
      - $ref: '#/parameters/qr_personal_account'
      - $ref: '#/parameters/qr_bank_name'
      - $ref: '#/parameters/qr_bic'
      - $ref: '#/parameters/qr_correspondent_account'
      - $ref: '#/parameters/qr_inn'
      - $ref: '#/parameters/qr_kpp'
      - $ref: '#/parameters/qr_payer_address'
    responses:
      200:
        description: DOCX файл с заполненными данными
        content:
          application/vnd.openxmlformats-officedocument.wordprocessingml.document:
            schema:
              type: string
              format: binary
      500:
        description: Ошибка генерации документа
    """
    try:
        replacements, payment_details, qr_width_mm = prepare_generation_inputs()
        docx_path, _, _ = build_doc(replacements, payment_details, qr_width_mm)
        return send_file(
            docx_path,
            download_name="document.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
        )
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/Document/GetPdfZip")
def get_pdf_zip():
    """Получить ZIP с PDF файлом
    ---
    tags:
      - Documents
    produces:
      - application/zip
    parameters:
      - $ref: '#/parameters/price'
      - $ref: '#/parameters/price_text'
      - $ref: '#/parameters/bill_date'
      - $ref: '#/parameters/invoiceDate'
      - $ref: '#/parameters/deal'
      - $ref: '#/parameters/service'
      - $ref: '#/parameters/city'
      - $ref: '#/parameters/lead_sum'
      - $ref: '#/parameters/lead_cost'
      - $ref: '#/parameters/revenue'
      - $ref: '#/parameters/email'
      - $ref: '#/parameters/phone'
      - $ref: '#/parameters/name'
      - $ref: '#/parameters/inn'
      - $ref: '#/parameters/companyName'
      - $ref: '#/parameters/qr_sum'
      - $ref: '#/parameters/qr_purpose'
      - $ref: '#/parameters/qr_width_mm'
      - $ref: '#/parameters/qr_name'
      - $ref: '#/parameters/qr_personal_account'
      - $ref: '#/parameters/qr_bank_name'
      - $ref: '#/parameters/qr_bic'
      - $ref: '#/parameters/qr_correspondent_account'
      - $ref: '#/parameters/qr_inn'
      - $ref: '#/parameters/qr_kpp'
      - $ref: '#/parameters/qr_payer_address'
    responses:
      200:
        description: ZIP архив с PDF
        content:
          application/zip:
            schema:
              type: string
              format: binary
      500:
        description: Ошибка генерации документа
    """
    try:
        replacements, payment_details, qr_width_mm = prepare_generation_inputs()
        _, pdf_path, _ = build_doc(replacements, payment_details, qr_width_mm)
        zip_buffer = zip_single_file(pdf_path, "document.pdf")
        return send_file(zip_buffer, download_name="document_pdf.zip", mimetype="application/zip", as_attachment=True)
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/Document/GetDocxZip")
def get_docx_zip():
    """Получить ZIP с DOCX файлом
    ---
    tags:
      - Documents
    produces:
      - application/zip
    parameters:
      - $ref: '#/parameters/price'
      - $ref: '#/parameters/price_text'
      - $ref: '#/parameters/bill_date'
      - $ref: '#/parameters/invoiceDate'
      - $ref: '#/parameters/deal'
      - $ref: '#/parameters/service'
      - $ref: '#/parameters/city'
      - $ref: '#/parameters/lead_sum'
      - $ref: '#/parameters/lead_cost'
      - $ref: '#/parameters/revenue'
      - $ref: '#/parameters/email'
      - $ref: '#/parameters/phone'
      - $ref: '#/parameters/name'
      - $ref: '#/parameters/inn'
      - $ref: '#/parameters/companyName'
      - $ref: '#/parameters/qr_sum'
      - $ref: '#/parameters/qr_purpose'
      - $ref: '#/parameters/qr_width_mm'
      - $ref: '#/parameters/qr_name'
      - $ref: '#/parameters/qr_personal_account'
      - $ref: '#/parameters/qr_bank_name'
      - $ref: '#/parameters/qr_bic'
      - $ref: '#/parameters/qr_correspondent_account'
      - $ref: '#/parameters/qr_inn'
      - $ref: '#/parameters/qr_kpp'
      - $ref: '#/parameters/qr_payer_address'
    responses:
      200:
        description: ZIP архив с DOCX
        content:
          application/zip:
            schema:
              type: string
              format: binary
      500:
        description: Ошибка генерации документа
    """
    try:
        replacements, payment_details, qr_width_mm = prepare_generation_inputs()
        docx_path, _, _ = build_doc(replacements, payment_details, qr_width_mm)
        zip_buffer = zip_single_file(docx_path, "document.docx")
        return send_file(zip_buffer, download_name="document_docx.zip", mimetype="application/zip", as_attachment=True)
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/Document/GetAllZip")
def get_all_zip():
    """Получить ZIP с DOCX, PDF и QR
    ---
    tags:
      - Documents
    produces:
      - application/zip
    parameters:
      - $ref: '#/parameters/price'
      - $ref: '#/parameters/price_text'
      - $ref: '#/parameters/bill_date'
      - $ref: '#/parameters/invoiceDate'
      - $ref: '#/parameters/deal'
      - $ref: '#/parameters/service'
      - $ref: '#/parameters/city'
      - $ref: '#/parameters/lead_sum'
      - $ref: '#/parameters/lead_cost'
      - $ref: '#/parameters/revenue'
      - $ref: '#/parameters/email'
      - $ref: '#/parameters/phone'
      - $ref: '#/parameters/name'
      - $ref: '#/parameters/inn'
      - $ref: '#/parameters/companyName'
      - $ref: '#/parameters/qr_sum'
      - $ref: '#/parameters/qr_purpose'
      - $ref: '#/parameters/qr_width_mm'
      - $ref: '#/parameters/qr_name'
      - $ref: '#/parameters/qr_personal_account'
      - $ref: '#/parameters/qr_bank_name'
      - $ref: '#/parameters/qr_bic'
      - $ref: '#/parameters/qr_correspondent_account'
      - $ref: '#/parameters/qr_inn'
      - $ref: '#/parameters/qr_kpp'
      - $ref: '#/parameters/qr_payer_address'
    responses:
      200:
        description: ZIP архив с документами и QR
        content:
          application/zip:
            schema:
              type: string
              format: binary
      500:
        description: Ошибка генерации документа
    """
    try:
        replacements, payment_details, qr_width_mm = prepare_generation_inputs()
        docx_path, pdf_path, qr_path = build_doc(replacements, payment_details, qr_width_mm)
        file_mappings = [
            (docx_path, "document.docx"),
            (pdf_path, "document.pdf"),
            (qr_path, "payment_qr.png")
        ]
        zip_buffer = zip_files(file_mappings)
        return send_file(zip_buffer, download_name="documents_full.zip", mimetype="application/zip", as_attachment=True)
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route("/Document/GetPaymentQr")
def get_payment_qr():
    """Получить PNG с банковским QR-кодом
    ---
    tags:
      - QR
    produces:
      - image/png
    parameters:
      - $ref: '#/parameters/price'
      - $ref: '#/parameters/price_text'
      - $ref: '#/parameters/bill_date'
      - $ref: '#/parameters/invoiceDate'
      - $ref: '#/parameters/deal'
      - $ref: '#/parameters/service'
      - $ref: '#/parameters/city'
      - $ref: '#/parameters/lead_sum'
      - $ref: '#/parameters/lead_cost'
      - $ref: '#/parameters/revenue'
      - $ref: '#/parameters/email'
      - $ref: '#/parameters/phone'
      - $ref: '#/parameters/name'
      - $ref: '#/parameters/inn'
      - $ref: '#/parameters/companyName'
      - $ref: '#/parameters/qr_sum'
      - $ref: '#/parameters/qr_purpose'
      - $ref: '#/parameters/qr_name'
      - $ref: '#/parameters/qr_personal_account'
      - $ref: '#/parameters/qr_bank_name'
      - $ref: '#/parameters/qr_bic'
      - $ref: '#/parameters/qr_correspondent_account'
      - $ref: '#/parameters/qr_inn'
      - $ref: '#/parameters/qr_kpp'
      - $ref: '#/parameters/qr_payer_address'
    responses:
      200:
        description: PNG файл с QR-кодом
        headers:
          X-Payment-QR-Payload-Base64:
            description: Base64-представление строки payload для QR
            schema:
              type: string
        content:
          image/png:
            schema:
              type: string
              format: binary
      400:
        description: QR-код не сформирован
      500:
        description: Ошибка генерации QR-кода
    """
    try:
        replacements = get_replacements()
        payment_details = get_payment_details(request.args, replacements)
        try:
            qr_payload, qr_path = generate_payment_qr_image(payment_details, str(uuid.uuid4()))
        except RuntimeError as dependency_error:
            return jsonify({"error": str(dependency_error)}), 500

        if not qr_payload or not qr_path or not os.path.exists(qr_path):
            return jsonify({"error": "Не удалось сформировать QR-код"}), 400

        with open(qr_path, "rb") as qr_file:
            buffer = BytesIO(qr_file.read())
        buffer.seek(0)

        response = send_file(
            buffer,
            download_name="payment_qr.png",
            mimetype="image/png",
            as_attachment=True,
        )
        payload_b64 = base64.b64encode(qr_payload.encode("utf-8")).decode("ascii")
        response.headers["X-Payment-QR-Payload-Base64"] = payload_b64

        try:
            os.remove(qr_path)
        except OSError:
            pass

        return response
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=12345, threaded=False)
